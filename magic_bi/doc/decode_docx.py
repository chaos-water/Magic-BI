import os.path

from loguru import logger
from docx import Document
from typing import List
import io

from magic_bi.model.text_embedding import TextEmbedding
from magic_bi.model.openai_adapter import OpenaiAdapter
from magic_bi.doc.doc_paragraph import DocParagraph


def decode_docx_by_bs(doc_bytes: bytes) -> str:
    from zipfile import ZipFile
    from io import BytesIO
    from bs4 import BeautifulSoup

    try:
        # 使用 BytesIO 将字节数据转换为文件对象
        with ZipFile(BytesIO(doc_bytes)) as document:
            output = ""
            xml = document.read("word/document.xml")
            word_obj = BeautifulSoup(xml.decode("utf-8"), "xml")  # 使用 "xml" 解析器
            texts = word_obj.find_all("w:t")
            for text in texts:
                output += text.get_text(strip=True)

            logger.debug("decode_docx_by_bytes suc, output length:%d" % len(output))
            return output  # 返回解析后的文本

    except Exception as e:
        logger.error("catch exception:%s" % str(e))
        return ""

def clean_text(input_content: str) -> str:
    import re
    return re.sub(r'\n+', '\n', input_content)

def decode_docx(file_bytes, max_chunk_size, text_embedding, openai_adapter):
    try:
        if docx_has_headings(file_bytes):
            doc_paragraph_list = decode_docx_with_heading(file_bytes, max_chunk_size, text_embedding, openai_adapter)
        else:
            doc_paragraph_list = parse_docx_without_headings(file_bytes, max_chunk_size, text_embedding, openai_adapter)

        full_content = get_content_from_docx(file_bytes)

        return doc_paragraph_list, full_content
    except Exception as e:
        full_content = decode_docx_by_bs(file_bytes)

        doc_paragraph = DocParagraph(text_embedding=text_embedding, openai_adapter=openai_adapter)
        doc_paragraph.content = full_content
        doc_paragraph.try_split_content()

        return [doc_paragraph], doc_paragraph.content

def docx_has_headings(file_bytes: bytes) -> bool:
    doc = Document(io.BytesIO(file_bytes))

    # 定义标题样式与级别的映射
    heading_styles = {
        'Heading 1',
        'Heading 2',
        'Heading 3',
        'Heading 4',
        'Heading 5',
        'Heading 6',
        'Heading 7',
        'Heading 8',
        'Heading 9'
    }

    # 遍历文档中的所有段落，检查样式是否为标题样式
    for para in doc.paragraphs:
        if para.style.name in heading_styles:
            return True

    return False


def parse_docx_without_headings(file_bytes: bytes, max_chunk_size: int, text_embedding: TextEmbedding, openai_adapter: OpenaiAdapter) -> DocParagraph:
    doc = Document(io.BytesIO(file_bytes))

    def extract_content(paragraphs):
        content = ""
        for para in paragraphs:
            content += para.text + "\n"
        return content

    paragraphs = doc.paragraphs
    # 获取整个文档的内容
    content = extract_content(paragraphs)
    if not content.strip():
        content = "无标题内容"

    doc_paragraph: DocParagraph = DocParagraph(text_embedding=text_embedding, openai_adapter=openai_adapter)
    doc_paragraph.content = content

    doc_paragraph.try_split_content()
    return [doc_paragraph]

def get_content_from_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    text = []
    for para in doc.paragraphs:
        text.append(para.text)

    content = '\n'.join(text)
    content = clean_text(content)
    return content

def decode_docx_with_heading(file_bytes: bytes, max_chunk_size: int, text_embedding: TextEmbedding, openai_adapter: OpenaiAdapter) -> List:
    doc = Document(io.BytesIO(file_bytes))

    # 定义标题样式与级别的映射
    heading_styles = {
        'Heading 1': 1,
        'Heading 2': 2,
        'Heading 3': 3,
        'Heading 4': 4,
        'Heading 5': 5,
        'Heading 6': 6,
        'Heading 7': 7,
        'Heading 8': 8,
        'Heading 9': 9
    }

    def add_sub_paragraphs(start_idx, paragraphs, current_level):
        doc_paragraph_list = []
        i = start_idx

        while i < len(paragraphs):
            para = paragraphs[i]
            style = para.style.name

            if style in heading_styles:
                heading_level = heading_styles[style]

                if heading_level == current_level:
                    doc_paragraph = DocParagraph(text_embedding=text_embedding, openai_adapter=openai_adapter, max_chunk_size=max_chunk_size)
                    doc_paragraph.title = para.text

                    # 处理下一个段落
                    i += 1
                    while i < len(paragraphs):
                        next_para = paragraphs[i]
                        next_style = next_para.style.name

                        if next_style in heading_styles:
                            next_level = heading_styles[next_style]

                            if next_level > current_level:
                                # 递归处理子级标题
                                sub_paragraphs, i = add_sub_paragraphs(i, paragraphs, next_level)
                                doc_paragraph.sub_paragraphs.extend(sub_paragraphs)
                            else:
                                # 如果遇到同级或上级标题，停止添加内容
                                break
                        else:
                            # 添加当前段落的内容
                            doc_paragraph.content += next_para.text + "\n"
                            i += 1

                    doc_paragraph.try_split_content()
                    # doc_paragraph_list.append(doc_paragraph.to_json())
                    doc_paragraph_list.append(doc_paragraph)
                elif heading_level < current_level:
                    # 如果遇到比当前级别低的标题，返回上一级处理
                    return doc_paragraph_list, i
                else:
                    # 遇到更高级别的标题，递归处理
                    sub_paragraphs, i = add_sub_paragraphs(i, paragraphs, heading_level)
                    doc_paragraph_list.extend(sub_paragraphs)
            else:
                # 跳过非标题段落
                i += 1

        return doc_paragraph_list, i

    # 解析文档中的段落
    paragraphs = doc.paragraphs
    content, _ = add_sub_paragraphs(0, paragraphs, 1)

    return content
